"""Markdown → python-docx converter for paper analysis reports.

Handles the most common LLM-generated markdown patterns:
  - ATX headings (# / ## / ### / ####)
  - Inline **bold**, *italic*, ***bold-italic***, `code`
  - Fenced code blocks (``` ... ```)
  - Unordered lists (- / * / +)
  - Ordered lists (1. 2. …)
  - Block-quotes (> …)
  - Horizontal rules (--- / *** / ___)
  - Plain paragraphs
"""

from __future__ import annotations

import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# ── Inline formatting ─────────────────────────────────────────────────────────

_INLINE_RE = re.compile(
    r'(\*\*\*.+?\*\*\*'   # ***bold italic***
    r'|\*\*.+?\*\*'        # **bold**
    r'|\*.+?\*'            # *italic*
    r'|`.+?`)',            # `inline code`
    re.DOTALL,
)


def _add_inline(paragraph, text: str) -> None:
    """Split *text* on bold/italic/code markers and add formatted runs."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


# ── Block-level markdown parser ───────────────────────────────────────────────

def _add_markdown(doc: Document, markdown: str, heading_offset: int = 0) -> None:
    """Parse *markdown* and append styled content to *doc*.

    Ordered-list numbering uses a per-call counter that resets whenever
    non-ordered-list content (headings, paragraphs, blockquotes, blank
    lines) appears, so each logical list group starts from 1.

    Args:
        heading_offset: Shift all heading levels by this amount (capped at 4).
                        Use 0 for single-paper export (angle = H1, content starts at H2).
                        Use 1 for batch export (angle = H2, content starts at H3).
    """
    if not markdown or not markdown.strip():
        return

    lines = markdown.split('\n')
    in_code = False
    code_lines: list[str] = []
    # Manual ordered-list counter — resets at any non-ordered-list line
    ordered_counter = 0
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Fenced code block ────────────────────────────
        if stripped.startswith('```'):
            ordered_counter = 0          # code fence breaks a list group
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── ATX headings ─────────────────────────────────
        # Base levels: #### → 4, ### → 3, ## → 2, # → 2 (downgraded since
        # we are already inside a higher-level section heading).
        # heading_offset shifts all levels further down (capped at 4).
        _h = lambda base: min(4, base + heading_offset)
        if stripped.startswith('#### '):
            ordered_counter = 0
            doc.add_heading(stripped[5:].strip(), level=_h(4))
        elif stripped.startswith('### '):
            ordered_counter = 0
            doc.add_heading(stripped[4:].strip(), level=_h(3))
        elif stripped.startswith('## '):
            ordered_counter = 0
            doc.add_heading(stripped[3:].strip(), level=_h(2))
        elif stripped.startswith('# '):
            ordered_counter = 0
            # Downgrade: content # is treated same as ## (we're inside a section)
            doc.add_heading(stripped[2:].strip(), level=_h(2))

        # ── Block-quote ──────────────────────────────────
        elif stripped.startswith('>'):
            ordered_counter = 0
            text = re.sub(r'^>+\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            _add_inline(p, text)

        # ── Horizontal rule ──────────────────────────────
        elif re.match(r'^[-*_]{3,}$', stripped):
            ordered_counter = 0
            p = doc.add_paragraph()
            p.add_run('─' * 42)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        # ── Unordered list ───────────────────────────────
        elif re.match(r'^[*\-+] ', stripped):
            ordered_counter = 0          # bullet list breaks an ordered group
            text = re.sub(r'^[*\-+] ', '', stripped)
            indent_lvl = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(style='List Bullet')
            if indent_lvl > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (1 + indent_lvl))
            _add_inline(p, text)

        # ── Ordered list (manual counter, resets per group) ──
        elif re.match(r'^\d+\. ', stripped):
            ordered_counter += 1
            text = re.sub(r'^\d+\.\s*', '', stripped)
            indent_lvl = (len(line) - len(line.lstrip())) // 2
            # Use a Normal paragraph with hanging indent so numbering is
            # controlled entirely by ordered_counter, not Word's list engine.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3 + 0.25 * indent_lvl)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f'{ordered_counter}.\u2002')  # en-space after dot
            num_run.bold = False
            _add_inline(p, text)

        # ── Empty line ───────────────────────────────────
        elif not stripped:
            ordered_counter = 0          # blank line ends a list group
            # Paragraph spacing handles visual separation

        # ── Regular paragraph ─────────────────────────────
        else:
            ordered_counter = 0
            p = doc.add_paragraph()
            _add_inline(p, stripped)

        i += 1


# ── Public API ────────────────────────────────────────────────────────────────

def build_docx(
    paper_title: str,
    angles: list[dict],
    final_report: Optional[str] = None,
) -> Document:
    """Build a Word document from paper analysis results.

    Args:
        paper_title: The paper's display title.
        angles: List of ``{'title': str, 'content': str}`` dicts, in order.
        final_report: Optional final-report markdown (appended on last page).

    Returns:
        A :class:`docx.Document` ready for ``.save()``.
    """
    doc = Document()

    # ── Cover page ────────────────────────────────────────
    title_p = doc.add_heading('论文分析报告', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(paper_title)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(datetime.now().strftime('%Y年%m月%d日'))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ── Angle sections ────────────────────────────────────
    valid_angles = [a for a in angles if a.get('content', '').strip()]
    for idx, angle in enumerate(valid_angles):
        doc.add_heading(angle['title'], level=1)
        _add_markdown(doc, angle['content'])
        needs_break = idx < len(valid_angles) - 1 or bool(final_report and final_report.strip())
        if needs_break:
            doc.add_page_break()

    # ── Final report ──────────────────────────────────────
    if final_report and final_report.strip():
        doc.add_heading('综合报告', level=1)
        _add_markdown(doc, final_report)

    return doc


def build_batch_docx(papers: list[dict]) -> Document:
    """Build a single Word document that contains analyses of multiple papers.

    Document structure:
        Cover page (title, paper count, date)
        ── Paper 1 title  [Heading 1]
            ── Angle 1    [Heading 2]
               content…
            ── Angle 2    [Heading 2]
               content…
            ── 综合报告   [Heading 2]  (if present)
        ── page break ──
        ── Paper 2 title  [Heading 1]
            …

    Args:
        papers: List of dicts, each with keys:
                  paper_title (str), angles (list[{title, content}]),
                  final_report (str | None).

    Returns:
        A :class:`docx.Document` ready for ``.save()``.
    """
    doc = Document()

    # ── Cover page ────────────────────────────────────────
    title_p = doc.add_heading('多篇论文分析汇总报告', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    count_p = doc.add_paragraph()
    count_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count_run = count_p.add_run(f'共 {len(papers)} 篇论文')
    count_run.font.size = Pt(13)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(datetime.now().strftime('%Y年%m月%d日'))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ── Per-paper sections ────────────────────────────────
    valid_papers = [p for p in papers if p.get('paper_title', '').strip()]
    for p_idx, paper in enumerate(valid_papers):
        # Paper title → H1  (major section boundary)
        doc.add_heading(paper['paper_title'], level=1)

        # Each angle → H2 ; content headings shifted +1 (start at H3)
        valid_angles = [a for a in paper.get('angles', []) if a.get('content', '').strip()]
        for angle in valid_angles:
            doc.add_heading(angle['title'], level=2)
            _add_markdown(doc, angle['content'], heading_offset=1)

        # Final report → H2 (if present)
        final = paper.get('final_report') or ''
        if final.strip():
            doc.add_heading('综合报告', level=2)
            _add_markdown(doc, final, heading_offset=1)

        # Page break between papers (not after the last one)
        if p_idx < len(valid_papers) - 1:
            doc.add_page_break()

    return doc
